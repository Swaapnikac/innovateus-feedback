"""Generate AWS Architecture .docx document"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

BLUE = RGBColor(0x12, 0x4D, 0x8F)
GOLD = RGBColor(0xC5, 0x8B, 0x00)
GRAY = RGBColor(0x66, 0x66, 0x66)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLUE

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Medium Shading 1 Accent 1'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def bold_para(bold_text, normal_text):
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True
    p.add_run(normal_text)

# ===== TITLE =====
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('InnovateUS Feedback Platform')
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AWS Serverless Architecture')
r.font.size = Pt(18); r.font.color.rgb = GOLD

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Lambda Functions | DynamoDB | API Gateway | CloudFront\nDetailed architecture with flows, costs, and security')
r.font.size = Pt(11); r.font.color.rgb = GRAY

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('April 2026')
r.font.size = Pt(12); r.font.color.rgb = GRAY

doc.add_page_break()

# ===== 1. OVERVIEW =====
heading('1. Architecture Overview')

doc.add_paragraph(
    'The InnovateUS Feedback Platform uses a fully serverless AWS architecture. '
    'The frontend is a statically-exported Next.js app hosted on S3 + CloudFront. '
    'The backend is split into 7 Lambda functions behind an API Gateway HTTP API, '
    'each serving a specific domain. Data is stored in 2 DynamoDB tables with on-demand billing.'
)

heading('High-Level Flow', 2)
doc.add_paragraph(
    'Users (Browser)\n'
    '    |\n'
    '    v\n'
    'CloudFront CDN --> S3 (Next.js Static Frontend)\n'
    '    |\n'
    '    v\n'
    'API Gateway (HTTP API)\n'
    '    |\n'
    '    +-- /v1/survey/*         --> Lambda 1: Survey Service\n'
    '    +-- /v1/submissions/*    --> Lambda 2: Submission Service\n'
    '    +-- /v1/ai/*             --> Lambda 3: AI Service\n'
    '    +-- /v1/transcribe       --> Lambda 4: Transcribe Service\n'
    '    +-- /v1/admin/*          --> Lambda 5: Admin Service\n'
    '    +-- /v1/admin/editor/*   --> Lambda 6: Editor Service\n'
    '    +-- /v1/admin/export/*   --> Lambda 7: Export Service\n'
    '    |\n'
    '    v\n'
    'DynamoDB (2 tables) + OpenAI API + Qualtrics/JotForm APIs'
)

doc.add_page_break()

# ===== 2. AWS SERVICES =====
heading('2. AWS Services Used')

table(
    ['Service', 'Purpose', 'Free Tier'],
    [
        ['API Gateway (HTTP API)', 'Route requests to Lambda functions', '1M requests/month'],
        ['Lambda (x7)', 'Run backend application logic', '1M requests + 400K GB-sec/month'],
        ['DynamoDB (x2 tables)', 'NoSQL data storage', '25 GB + 25 RCU/WCU'],
        ['S3', 'Host Next.js static site', '5 GB storage'],
        ['CloudFront', 'CDN + HTTPS', '1 TB transfer/month'],
        ['IAM', 'Roles and permissions', 'Free'],
        ['CloudWatch', 'Logs and monitoring', '5 GB logs/month'],
        ['ACM', 'SSL/TLS certificates', 'Free'],
    ]
)

doc.add_page_break()

# ===== 3. LAMBDA FUNCTIONS =====
heading('3. Lambda Functions (Detailed)')

doc.add_paragraph('The backend is organized into 7 Lambda functions, each handling a specific domain.')

heading('Lambda Summary Table', 2)
table(
    ['#', 'Name', 'Routes', 'Memory', 'Timeout', 'External APIs'],
    [
        ['1', 'Survey Service', 'GET /v1/survey/{id}', '256 MB', '10s', 'None'],
        ['2', 'Submission Service', 'POST /v1/submissions/*', '512 MB', '60s', 'OpenAI GPT-4o, Qualtrics'],
        ['3', 'AI Service', 'POST /v1/ai/*', '256 MB', '30s', 'OpenAI GPT-4o-mini'],
        ['4', 'Transcribe Service', 'POST /v1/transcribe', '512 MB', '30s', 'OpenAI Whisper'],
        ['5', 'Admin Service', '/v1/admin/* (excl. editor, export)', '512 MB', '120s', 'Qualtrics, JotForm'],
        ['6', 'Editor Service', '/v1/admin/editor/*', '256 MB', '15s', 'None'],
        ['7', 'Export Service', '/v1/admin/export/*', '1024 MB', '60s', 'None'],
    ]
)

# Lambda 1
heading('3.1 Lambda 1: Survey Service', 2)
doc.add_paragraph('Serves survey configuration to end users.')
table(['Property', 'Value'], [
    ['Route', 'GET /v1/survey/{cohort_id}'],
    ['Memory', '256 MB'], ['Timeout', '10 seconds'],
    ['DynamoDB', 'READ only (surveys table)'],
    ['Auth', 'None (public)'],
])
doc.add_paragraph()
bullet('Reads cohort metadata from DynamoDB surveys table')
bullet('Gets the active survey_config')
bullet('Randomizes questions within groups (respects conditional dependencies)')
bullet('Returns JSON with Cache-Control: no-store')

# Lambda 2
heading('3.2 Lambda 2: Submission Service', 2)
doc.add_paragraph('Handles the full survey submission lifecycle.')
table(['Property', 'Value'], [
    ['Routes', 'POST /v1/submissions/start\nPOST /v1/submissions/{id}/answer\nPOST /v1/submissions/{id}/complete'],
    ['Memory', '512 MB'], ['Timeout', '60 seconds'],
    ['DynamoDB', 'READ + WRITE (both tables + GSI)'],
    ['External APIs', 'OpenAI GPT-4o (extraction), Qualtrics (auto-sync)'],
    ['Auth', 'None (public)'],
])

doc.add_paragraph()
heading('Start Submission Flow:', 3)
bullet('1. Get cohort metadata (max_submissions_per_ip)')
bullet('2. Hash client IP (SHA256 with salt)')
bullet('3. Query IpHashIndex GSI: check completed count for this IP')
bullet('4. If >= limit: return 429 "Already submitted"')
bullet('5. Check for existing in-progress submission -> return if found')
bullet('6. Create new submission in DynamoDB, return submission_id')

heading('Save Answer Flow:', 3)
bullet('1. Find submission by scanning for submission_id')
bullet('2. Upsert answer in embedded answers[] array')
bullet('3. Update DynamoDB item')

heading('Complete Submission Flow:', 3)
bullet('1. Find submission from DynamoDB')
bullet('2. Gather all answers (including followup answers)')
bullet('3. Call OpenAI GPT-4o: extract_structured()')
bullet('   -> Extracts: themes, barriers, enablers, success stories, workflows')
bullet('4. On AI error: gracefully store empty extraction')
bullet('5. Update: status=completed, completed_at, time_to_complete_sec, extraction')
bullet('6. Attempt Qualtrics sync (non-blocking, logs warning on failure)')
bullet('7. Set duplicate-prevention cookie (30 day expiry)')
bullet('8. Return { status: "completed", extraction }')

# Lambda 3
heading('3.3 Lambda 3: AI Service', 2)
doc.add_paragraph('Real-time AI analysis during survey.')
table(['Property', 'Value'], [
    ['Routes', 'POST /v1/ai/vagueness\nPOST /v1/ai/followups'],
    ['Memory', '256 MB'], ['Timeout', '30 seconds'],
    ['DynamoDB', 'None'], ['External APIs', 'OpenAI GPT-4o-mini'],
])
doc.add_paragraph()
bold_para('Vagueness Detection: ', 'Input: question + answer. Uses GPT-4o-mini (temp=0.1). Returns: { is_vague, reason, missing_info_types[] }')
bold_para('Follow-up Generation: ', 'Input: question + answer + missing types. Uses GPT-4o-mini (temp=0.3). Returns: { followups: ["Q1?", "Q2?"] } (max 2)')

# Lambda 4
heading('3.4 Lambda 4: Transcribe Service', 2)
doc.add_paragraph('Converts voice recordings to text.')
table(['Property', 'Value'], [
    ['Route', 'POST /v1/transcribe'],
    ['Memory', '512 MB'], ['Timeout', '30 seconds'],
    ['Input', 'multipart/form-data with WebM audio file'],
    ['External APIs', 'OpenAI Whisper (whisper-1)'],
    ['Privacy', 'Audio is never stored - only transcript text is kept'],
])

# Lambda 5
heading('3.5 Lambda 5: Admin Service', 2)
doc.add_paragraph('Dashboard, metrics, cohort management, and pipeline integrations.')
table(['Route', 'Method', 'Purpose'], [
    ['/v1/admin/login', 'POST', 'Authenticate admin'],
    ['/v1/admin/cohorts', 'GET', 'List all cohorts/programs'],
    ['/v1/admin/cohorts', 'POST', 'Create new cohort'],
    ['/v1/admin/cohorts/{id}/settings', 'POST', 'Update max submissions per IP'],
    ['/v1/admin/metrics', 'GET', 'Aggregate metrics with filters'],
    ['/v1/admin/responses', 'GET', 'Paginated submissions with answers + extraction'],
    ['/v1/admin/responses', 'DELETE', 'Bulk delete submissions'],
    ['/v1/admin/qualtrics/status', 'GET', 'Check Qualtrics configuration'],
    ['/v1/admin/qualtrics/sync/{id}', 'POST', 'Sync one submission to Qualtrics'],
    ['/v1/admin/qualtrics/sync-all', 'POST', 'Bulk sync all to Qualtrics'],
    ['/v1/admin/jotform/status', 'GET', 'Check JotForm configuration'],
    ['/v1/admin/jotform/sync/{id}', 'POST', 'Sync one submission to JotForm'],
    ['/v1/admin/jotform/sync-all', 'POST', 'Bulk sync all to JotForm'],
])
doc.add_paragraph()
table(['Property', 'Value'], [
    ['Memory', '512 MB'], ['Timeout', '120 seconds (bulk sync)'],
    ['DynamoDB', 'READ + WRITE (both tables)'],
    ['External APIs', 'Qualtrics Response Import API, JotForm Submissions API'],
])

doc.add_paragraph()
heading('Metrics Computed:', 3)
bullet('total_submissions, completed_submissions, completion_rate')
bullet('avg_time_to_complete_sec')
bullet('avg_recommend_score (from q1_recommend)')
bullet('confidence_distribution (from q2_confidence)')
bullet('vagueness_rate (% of open-ended answers flagged vague)')

# Lambda 6
heading('3.6 Lambda 6: Editor Service', 2)
doc.add_paragraph('Survey question editor with automatic version control.')
table(['Route', 'Method', 'Purpose'], [
    ['/v1/admin/editor/login', 'POST', 'Authenticate editor'],
    ['/v1/admin/editor/cohorts', 'GET', 'List cohorts'],
    ['/v1/admin/editor/survey/{id}', 'GET', 'Get current config for editing'],
    ['/v1/admin/editor/survey/{id}', 'PUT', 'Save changes (creates version)'],
    ['/v1/admin/editor/survey/{id}/versions', 'GET', 'List version history'],
    ['/v1/admin/editor/survey/{id}/versions/{v}', 'GET', 'View specific version'],
    ['/v1/admin/editor/survey/{id}/versions/{v}/restore', 'POST', 'Restore old version'],
])
doc.add_paragraph()
bullet('Auto-generates change summary (e.g., "Changed text: q3; Added: q10_new")')
bullet('Detects "no changes" via deep JSON comparison')
bullet('Each save creates VERSION#vN in DynamoDB')
bullet('Restore is non-destructive: creates new version from old config')

# Lambda 7
heading('3.7 Lambda 7: Export Service', 2)
doc.add_paragraph('Generates downloadable reports in 4 formats.')
table(['Route', 'Format', 'Description'], [
    ['/v1/admin/export/raw.csv', 'CSV', 'One row per answer with all followups'],
    ['/v1/admin/export/structured.csv', 'CSV', 'One row per respondent'],
    ['/v1/admin/export/summary.pdf', 'PDF', 'Branded report: metrics + themes + barriers + stories'],
    ['/v1/admin/export/summary.pptx', 'PPTX', '6-slide presentation for stakeholders'],
])
doc.add_paragraph()
table(['Property', 'Value'], [
    ['Memory', '1024 MB (PDF/PPTX generation)'],
    ['Timeout', '60 seconds'],
    ['Libraries', 'reportlab (PDF), python-pptx (PPTX), pandas (CSV)'],
])

doc.add_page_break()

# ===== 4. DYNAMODB =====
heading('4. DynamoDB Table Design')

heading('Table 1: innovateus-surveys', 2)
doc.add_paragraph('Stores cohort metadata and survey configuration versions.')
table(['pk (Partition Key)', 'sk (Sort Key)', 'Key Attributes'], [
    ['COHORT#{uuid}', 'METADATA', 'name, course_name, survey_config, active_version, max_submissions_per_ip, created_at'],
    ['COHORT#{uuid}', 'VERSION#v1', 'config, change_summary, created_by, created_at'],
    ['COHORT#{uuid}', 'VERSION#v2', 'config, change_summary, created_by, created_at'],
])
doc.add_paragraph('Billing: PAY_PER_REQUEST. No GSIs needed.')

heading('Table 2: innovateus-submissions', 2)
doc.add_paragraph('Stores submissions with embedded answers and AI extraction.')
table(['pk', 'sk', 'Key Attributes'], [
    ['COHORT#{uuid}', 'SUB#{sub_uuid}', 'submission_id, status, created_at, completed_at, ip_hash, answers[], extraction, survey_version'],
])
doc.add_paragraph()
doc.add_paragraph('GSI: IpHashIndex (pk=ip_hash, sk=created_at) - for fast duplicate detection.')
doc.add_paragraph('Billing: PAY_PER_REQUEST.')

doc.add_paragraph()
heading('Embedded Answers Array:', 3)
bullet('question_id, question_type, answer_raw, input_mode')
bullet('is_vague, followups_asked, transcript')
bullet('followup_1 (question), followup_1_answer (response)')
bullet('followup_2 (question), followup_2_answer (response)')

heading('Extraction Object:', 3)
bullet('what_was_tried, planned_task_or_workflow, outcome_or_expected_outcome')
bullet('barriers[], enablers[], top_themes[]')
bullet('public_benefit, success_story_candidate')

doc.add_page_break()

# ===== 5. FLOWS =====
heading('5. End-to-End Request Flows')

heading('Flow 1: User Takes a Survey', 2)
bold_para('Step 1 - Open Survey: ', 'User clicks link -> CloudFront -> S3 serves consent page')
bold_para('Step 2 - Accept Consent: ', 'Lambda 1 fetches survey config, Lambda 2 creates submission')
bold_para('Step 3 - Answer Questions: ', 'For each question: Lambda 2 saves answer. If open-ended: Lambda 3 checks vagueness -> generates followups. If voice: Lambda 4 transcribes.')
bold_para('Step 4 - Complete: ', 'Lambda 2 runs GPT-4o extraction, syncs Qualtrics, sets cookie')
bold_para('Step 5 - Thank You: ', 'Frontend shows extraction summary')

heading('Flow 2: Admin Dashboard', 2)
bold_para('Step 1 - Login: ', 'Lambda 5 verifies password, returns JWT')
bold_para('Step 2 - Load: ', 'Parallel: cohorts + metrics + responses from Lambda 5')
bold_para('Step 3 - Export: ', 'Lambda 7 generates PDF/CSV/PPTX')
bold_para('Step 4 - Sync: ', 'Lambda 5 pushes to Qualtrics/JotForm')

heading('Flow 3: Editor Modifies Survey', 2)
bold_para('Step 1 - Login: ', 'Lambda 6 verifies editor password, returns JWT')
bold_para('Step 2 - Load: ', 'Lambda 6 returns cohorts + current config + version history')
bold_para('Step 3 - Save: ', 'Lambda 6 diffs configs, generates change summary, creates new version')
bold_para('Step 4 - Restore: ', 'Lambda 6 creates new version from old config')

doc.add_page_break()

# ===== 6. SECURITY =====
heading('6. Security Architecture')

table(['Layer', 'Measures'], [
    ['CloudFront', 'HTTPS only (TLS 1.2+), custom domain, ACM cert, optional WAF'],
    ['API Gateway', 'CORS per-origin, rate limiting, request validation'],
    ['Lambda', 'JWT auth (HS256, 24hr), role-based access (admin/editor)'],
    ['DynamoDB', 'Encryption at rest, IAM access, point-in-time recovery'],
    ['Data Privacy', 'IPs hashed (SHA256+salt), no PII, audio never stored'],
])

doc.add_page_break()

# ===== 7. IAM =====
heading('7. IAM Roles (Least Privilege)')

table(['Lambda', 'DynamoDB Actions', 'Tables'], [
    ['Survey', 'GetItem', 'surveys only'],
    ['Submission', 'GetItem, PutItem, UpdateItem, Query, Scan', 'Both + GSI'],
    ['AI', 'None', 'None'],
    ['Transcribe', 'None', 'None'],
    ['Admin', 'All CRUD + BatchWriteItem', 'Both + GSI'],
    ['Editor', 'GetItem, PutItem, UpdateItem, Query, Scan', 'surveys only'],
    ['Export', 'GetItem, Query, Scan', 'Both (read-only)'],
])

doc.add_page_break()

# ===== 8. COST =====
heading('8. Cost Estimates')
doc.add_paragraph('Based on ~100-500 submissions/month.')

heading('AWS Costs', 2)
table(['Component', 'Monthly Cost'], [
    ['Lambda', '$0 (free tier)'],
    ['DynamoDB', '$0 (free tier)'],
    ['API Gateway', '$0 (free tier)'],
    ['S3 + CloudFront', '$0-1'],
    ['AWS Total', '$0-1/month'],
])

heading('External API Costs', 2)
table(['Service', 'Monthly Cost'], [
    ['OpenAI GPT-4o (extraction)', '~$1-5'],
    ['OpenAI GPT-4o-mini (vagueness)', '~$0.50-2'],
    ['OpenAI Whisper (voice)', '~$0-2'],
    ['External Total', '~$2-9/month'],
])

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Total: ~$2-10/month (vs ~$7+/month on Render with PostgreSQL)')
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = BLUE

doc.add_page_break()

# ===== 9. DEPLOYMENT =====
heading('9. Deployment Strategy')

heading('Recommended: Start with Single Lambda', 2)
doc.add_paragraph(
    'Your codebase already has Mangum (FastAPI-to-Lambda adapter). '
    'Deploy the entire app as one Lambda behind API Gateway. '
    'Split into 7 Lambdas later when you need optimization.'
)

table(['Approach', 'Pros', 'Cons'], [
    ['Single Lambda', 'Simple, fast to deploy, one package', 'Cold starts affect all, shared memory'],
    ['7 Lambdas', 'Independent scaling, least-privilege, isolated', 'More config and deployment work'],
])

heading('CI/CD with GitHub Actions', 2)
bullet('On push to main: run tests')
bullet('Build Next.js -> upload to S3')
bullet('Package Lambda -> deploy via SAM/CDK')
bullet('Invalidate CloudFront cache')

# Save
doc.save('docs/InnovateUS_AWS_Architecture.docx')
print('Saved: docs/InnovateUS_AWS_Architecture.docx')
